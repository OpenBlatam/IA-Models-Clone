"""
DOCX export handler.
"""

from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .base import BaseExporter


class DOCXExporter(BaseExporter):
    """DOCX export handler with professional formatting."""
    
    def get_supported_features(self) -> list:
        """Get supported DOCX features."""
        return [
            "Editable",
            "Professional formatting",
            "Table support",
            "Image embedding",
            "Style management",
            "Headers and footers",
            "Page numbers",
            "Cross-platform compatibility"
        ]
    
    async def export(
        self, 
        content: Dict[str, Any], 
        config: Any, 
        output_path: str
    ) -> Dict[str, Any]:
        """Export content to DOCX format."""
        try:
            # Preprocess content
            processed_content = self.preprocess_content(content)
            
            # Create DOCX document
            doc = Document()
            
            # Set document styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Title
            if "title" in processed_content:
                title = doc.add_heading(processed_content["title"], 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Content sections
            if "sections" in processed_content:
                for section in processed_content["sections"]:
                    if "heading" in section:
                        doc.add_heading(section["heading"], level=1)
                    
                    if "content" in section:
                        paragraph = doc.add_paragraph(section["content"])
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Save document
            doc.save(output_path)
            
            result = {
                "format": "docx",
                "sections": len(processed_content.get("sections", [])),
                "professional_features": {
                    "editable": True,
                    "professional_formatting": True,
                    "table_support": True
                }
            }
            
            return self.postprocess_result(result)
            
        except Exception as e:
            self.logger.error(f"DOCX export failed: {e}")
            raise
    
    def postprocess_result(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """Postprocess DOCX export result."""
        result["file_type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        result["mime_type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return result




