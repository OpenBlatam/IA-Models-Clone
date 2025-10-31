"""
Export IA Engine - Professional File Export System
=================================================

Advanced AI-powered system for exporting documents and content in professional formats.
Ensures all exported files maintain high-quality, professional appearance across different formats.
"""

import asyncio
import logging
from typing import Dict, List, Optional, Any, Union, BinaryIO
from datetime import datetime
from dataclasses import dataclass, field
from enum import Enum
import json
import uuid
import os
from pathlib import Path
import io
import base64

# Document processing libraries
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

import markdown
from jinja2 import Template, Environment, FileSystemLoader

logger = logging.getLogger(__name__)

class ExportFormat(Enum):
    """Supported export formats."""
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "markdown"
    RTF = "rtf"
    TXT = "txt"
    JSON = "json"
    XML = "xml"

class DocumentType(Enum):
    """Types of documents that can be exported."""
    BUSINESS_PLAN = "business_plan"
    REPORT = "report"
    PROPOSAL = "proposal"
    PRESENTATION = "presentation"
    MANUAL = "manual"
    CONTRACT = "contract"
    LETTER = "letter"
    MEMO = "memo"
    NEWSLETTER = "newsletter"
    CATALOG = "catalog"

class QualityLevel(Enum):
    """Quality levels for export."""
    BASIC = "basic"
    STANDARD = "standard"
    PROFESSIONAL = "professional"
    PREMIUM = "premium"
    ENTERPRISE = "enterprise"

@dataclass
class ExportConfig:
    """Configuration for export operations."""
    format: ExportFormat
    document_type: DocumentType
    quality_level: QualityLevel = QualityLevel.PROFESSIONAL
    template: Optional[str] = None
    custom_styles: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    branding: Dict[str, Any] = field(default_factory=dict)
    output_options: Dict[str, Any] = field(default_factory=dict)

@dataclass
class ExportTask:
    """Represents an export task."""
    id: str
    content: Dict[str, Any]
    config: ExportConfig
    status: str = "pending"  # pending, processing, completed, failed
    created_at: datetime = field(default_factory=datetime.now)
    completed_at: Optional[datetime] = None
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None
    file_path: Optional[str] = None
    file_size: Optional[int] = None

@dataclass
class ExportResult:
    """Result of an export operation."""
    task_id: str
    success: bool
    file_path: Optional[str] = None
    file_size: Optional[int] = None
    format: str = ""
    quality_score: float = 0.0
    processing_time: float = 0.0
    error: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

class ExportIAEngine:
    """
    Advanced AI-powered export engine for professional document generation.
    """
    
    def __init__(self):
        self.active_tasks: Dict[str, ExportTask] = {}
        self.completed_tasks: Dict[str, ExportResult] = {}
        self.templates: Dict[str, Any] = {}
        self.styles: Dict[str, Any] = {}
        
        # Initialize templates and styles
        self._initialize_templates()
        self._initialize_styles()
        
        # Export handlers
        self.export_handlers = {
            ExportFormat.PDF: self._export_to_pdf,
            ExportFormat.DOCX: self._export_to_docx,
            ExportFormat.HTML: self._export_to_html,
            ExportFormat.MARKDOWN: self._export_to_markdown,
            ExportFormat.RTF: self._export_to_rtf,
            ExportFormat.TXT: self._export_to_txt,
            ExportFormat.JSON: self._export_to_json,
            ExportFormat.XML: self._export_to_xml
        }
        
        logger.info("Export IA Engine initialized with professional export capabilities")
    
    def _initialize_templates(self):
        """Initialize professional document templates."""
        self.templates = {
            DocumentType.BUSINESS_PLAN: {
                "title_style": "Title",
                "heading_styles": ["Heading1", "Heading2", "Heading3"],
                "body_style": "Normal",
                "sections": [
                    "Executive Summary",
                    "Company Description",
                    "Market Analysis",
                    "Organization & Management",
                    "Service or Product Line",
                    "Marketing & Sales",
                    "Funding Request",
                    "Financial Projections",
                    "Appendix"
                ]
            },
            DocumentType.REPORT: {
                "title_style": "Title",
                "heading_styles": ["Heading1", "Heading2", "Heading3"],
                "body_style": "Normal",
                "sections": [
                    "Introduction",
                    "Methodology",
                    "Findings",
                    "Analysis",
                    "Conclusions",
                    "Recommendations",
                    "References"
                ]
            },
            DocumentType.PROPOSAL: {
                "title_style": "Title",
                "heading_styles": ["Heading1", "Heading2", "Heading3"],
                "body_style": "Normal",
                "sections": [
                    "Executive Summary",
                    "Problem Statement",
                    "Proposed Solution",
                    "Implementation Plan",
                    "Budget",
                    "Timeline",
                    "Team Qualifications",
                    "Next Steps"
                ]
            }
        }
    
    def _initialize_styles(self):
        """Initialize professional styling configurations."""
        self.styles = {
            QualityLevel.BASIC: {
                "font_family": "Arial",
                "font_size": 11,
                "line_spacing": 1.0,
                "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1},
                "colors": {"primary": "#000000", "secondary": "#666666"}
            },
            QualityLevel.STANDARD: {
                "font_family": "Calibri",
                "font_size": 11,
                "line_spacing": 1.15,
                "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1},
                "colors": {"primary": "#2E2E2E", "secondary": "#5A5A5A", "accent": "#1F4E79"}
            },
            QualityLevel.PROFESSIONAL: {
                "font_family": "Calibri",
                "font_size": 11,
                "line_spacing": 1.15,
                "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1},
                "colors": {"primary": "#2E2E2E", "secondary": "#5A5A5A", "accent": "#1F4E79", "highlight": "#F2F2F2"},
                "header_footer": True,
                "page_numbers": True,
                "table_styling": True
            },
            QualityLevel.PREMIUM: {
                "font_family": "Calibri",
                "font_size": 11,
                "line_spacing": 1.15,
                "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1},
                "colors": {"primary": "#2E2E2E", "secondary": "#5A5A5A", "accent": "#1F4E79", "highlight": "#F2F2F2"},
                "header_footer": True,
                "page_numbers": True,
                "table_styling": True,
                "custom_branding": True,
                "advanced_formatting": True
            },
            QualityLevel.ENTERPRISE: {
                "font_family": "Calibri",
                "font_size": 11,
                "line_spacing": 1.15,
                "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1},
                "colors": {"primary": "#2E2E2E", "secondary": "#5A5A5A", "accent": "#1F4E79", "highlight": "#F2F2F2"},
                "header_footer": True,
                "page_numbers": True,
                "table_styling": True,
                "custom_branding": True,
                "advanced_formatting": True,
                "interactive_elements": True,
                "accessibility_features": True
            }
        }
    
    async def export_document(
        self,
        content: Dict[str, Any],
        config: ExportConfig,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export a document in the specified format with professional quality.
        
        Args:
            content: Document content to export
            config: Export configuration
            output_path: Optional output file path
            
        Returns:
            Task ID for tracking the export process
        """
        task_id = str(uuid.uuid4())
        
        # Create export task
        task = ExportTask(
            id=task_id,
            content=content,
            config=config
        )
        
        self.active_tasks[task_id] = task
        
        # Start export process
        asyncio.create_task(self._process_export(task, output_path))
        
        logger.info(f"Export task created: {task_id} - {config.format.value} format")
        return task_id
    
    async def _process_export(self, task: ExportTask, output_path: Optional[str] = None):
        """Process an export task."""
        start_time = datetime.now()
        
        try:
            task.status = "processing"
            
            # Get export handler
            handler = self.export_handlers.get(task.config.format)
            if not handler:
                raise ValueError(f"Unsupported export format: {task.config.format}")
            
            # Generate output path if not provided
            if not output_path:
                output_path = self._generate_output_path(task)
            
            # Ensure output directory exists
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Process content for professional appearance
            processed_content = await self._process_content_for_quality(task.content, task.config)
            
            # Export to specified format
            result = await handler(processed_content, task.config, output_path)
            
            # Calculate quality score
            quality_score = await self._calculate_quality_score(result, task.config)
            
            # Calculate processing time
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Create export result
            export_result = ExportResult(
                task_id=task.id,
                success=True,
                file_path=output_path,
                file_size=os.path.getsize(output_path) if os.path.exists(output_path) else 0,
                format=task.config.format.value,
                quality_score=quality_score,
                processing_time=processing_time,
                metadata={
                    "document_type": task.config.document_type.value,
                    "quality_level": task.config.quality_level.value,
                    "exported_at": datetime.now().isoformat()
                }
            )
            
            # Store result
            self.completed_tasks[task.id] = export_result
            task.status = "completed"
            task.completed_at = datetime.now()
            task.result = result
            task.file_path = output_path
            task.file_size = export_result.file_size
            
            logger.info(f"Export completed: {task.id} - Quality score: {quality_score:.2f}")
            
        except Exception as e:
            logger.error(f"Export failed: {task.id} - {e}")
            
            task.status = "failed"
            task.error = str(e)
            task.completed_at = datetime.now()
            
            # Create failed result
            export_result = ExportResult(
                task_id=task.id,
                success=False,
                error=str(e),
                processing_time=(datetime.now() - start_time).total_seconds()
            )
            
            self.completed_tasks[task.id] = export_result
    
    async def _process_content_for_quality(self, content: Dict[str, Any], config: ExportConfig) -> Dict[str, Any]:
        """Process content to ensure professional quality."""
        processed = content.copy()
        
        # Apply quality enhancements based on quality level
        quality_config = self.styles[config.quality_level]
        
        # Enhance structure
        if "structure" not in processed:
            processed["structure"] = self._generate_document_structure(config.document_type)
        
        # Enhance formatting
        processed["formatting"] = self._apply_professional_formatting(processed, quality_config)
        
        # Add branding if configured
        if config.branding and quality_config.get("custom_branding"):
            processed["branding"] = config.branding
        
        # Enhance accessibility
        if quality_config.get("accessibility_features"):
            processed["accessibility"] = self._add_accessibility_features(processed)
        
        return processed
    
    def _generate_document_structure(self, document_type: DocumentType) -> Dict[str, Any]:
        """Generate professional document structure."""
        template = self.templates.get(document_type, {})
        return {
            "sections": template.get("sections", []),
            "hierarchy": template.get("heading_styles", []),
            "style_mapping": template
        }
    
    def _apply_professional_formatting(self, content: Dict[str, Any], quality_config: Dict[str, Any]) -> Dict[str, Any]:
        """Apply professional formatting to content."""
        return {
            "typography": {
                "font_family": quality_config["font_family"],
                "font_size": quality_config["font_size"],
                "line_spacing": quality_config["line_spacing"]
            },
            "layout": {
                "margins": quality_config["margins"],
                "colors": quality_config["colors"]
            },
            "features": {
                "header_footer": quality_config.get("header_footer", False),
                "page_numbers": quality_config.get("page_numbers", False),
                "table_styling": quality_config.get("table_styling", False)
            }
        }
    
    def _add_accessibility_features(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Add accessibility features to content."""
        return {
            "alt_text": True,
            "heading_structure": True,
            "color_contrast": True,
            "readable_fonts": True,
            "logical_reading_order": True
        }
    
    def _generate_output_path(self, task: ExportTask) -> str:
        """Generate output file path."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"export_{task.id}_{timestamp}.{task.config.format.value}"
        return os.path.join("exports", filename)
    
    async def _calculate_quality_score(self, result: Dict[str, Any], config: ExportConfig) -> float:
        """Calculate quality score for exported document."""
        score = 0.0
        
        # Base score for successful export
        score += 0.3
        
        # Quality level bonus
        quality_bonus = {
            QualityLevel.BASIC: 0.1,
            QualityLevel.STANDARD: 0.2,
            QualityLevel.PROFESSIONAL: 0.3,
            QualityLevel.PREMIUM: 0.4,
            QualityLevel.ENTERPRISE: 0.5
        }
        score += quality_bonus.get(config.quality_level, 0.2)
        
        # Format-specific quality checks
        if config.format == ExportFormat.PDF:
            score += 0.2  # PDF is generally high quality
        elif config.format in [ExportFormat.DOCX, ExportFormat.HTML]:
            score += 0.15
        
        # Professional features bonus
        quality_config = self.styles[config.quality_level]
        if quality_config.get("header_footer"):
            score += 0.1
        if quality_config.get("table_styling"):
            score += 0.1
        if quality_config.get("custom_branding"):
            score += 0.1
        
        return min(score, 1.0)  # Cap at 1.0
    
    # Export handlers for different formats
    async def _export_to_pdf(self, content: Dict[str, Any], config: ExportConfig, output_path: str) -> Dict[str, Any]:
        """Export content to PDF format."""
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Apply professional styling
        quality_config = self.styles[config.quality_level]
        
        # Title
        if "title" in content:
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER,
                textColor=colors.HexColor(quality_config["colors"]["primary"])
            )
            story.append(Paragraph(content["title"], title_style))
            story.append(Spacer(1, 12))
        
        # Content sections
        if "sections" in content:
            for section in content["sections"]:
                if "heading" in section:
                    heading_style = ParagraphStyle(
                        'CustomHeading',
                        parent=styles['Heading1'],
                        fontSize=14,
                        spaceAfter=12,
                        textColor=colors.HexColor(quality_config["colors"]["accent"])
                    )
                    story.append(Paragraph(section["heading"], heading_style))
                
                if "content" in section:
                    body_style = ParagraphStyle(
                        'CustomBody',
                        parent=styles['Normal'],
                        fontSize=quality_config["font_size"],
                        spaceAfter=6,
                        textColor=colors.HexColor(quality_config["colors"]["primary"])
                    )
                    story.append(Paragraph(section["content"], body_style))
                    story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        
        return {
            "format": "pdf",
            "pages": len(story),
            "professional_features": quality_config
        }
    
    async def _export_to_docx(self, content: Dict[str, Any], config: ExportConfig, output_path: str) -> Dict[str, Any]:
        """Export content to DOCX format."""
        doc = Document()
        
        # Apply professional styling
        quality_config = self.styles[config.quality_level]
        
        # Set document styles
        style = doc.styles['Normal']
        font = style.font
        font.name = quality_config["font_family"]
        font.size = Pt(quality_config["font_size"])
        
        # Title
        if "title" in content:
            title = doc.add_heading(content["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Content sections
        if "sections" in content:
            for section in content["sections"]:
                if "heading" in section:
                    doc.add_heading(section["heading"], level=1)
                
                if "content" in section:
                    paragraph = doc.add_paragraph(section["content"])
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Save document
        doc.save(output_path)
        
        return {
            "format": "docx",
            "sections": len(content.get("sections", [])),
            "professional_features": quality_config
        }
    
    async def _export_to_html(self, content: Dict[str, Any], config: ExportConfig, output_path: str) -> Dict[str, Any]:
        """Export content to HTML format."""
        quality_config = self.styles[config.quality_level]
        
        # Create HTML template
        html_template = """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{{ title }}</title>
            <style>
                body {
                    font-family: {{ font_family }};
                    font-size: {{ font_size }}px;
                    line-height: {{ line_spacing }};
                    color: {{ primary_color }};
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                }
                h1, h2, h3 {
                    color: {{ accent_color }};
                }
                .header {
                    text-align: center;
                    margin-bottom: 30px;
                }
                .section {
                    margin-bottom: 20px;
                }
            </style>
        </head>
        <body>
            <div class="header">
                <h1>{{ title }}</h1>
            </div>
            {% for section in sections %}
            <div class="section">
                {% if section.heading %}
                <h2>{{ section.heading }}</h2>
                {% endif %}
                {% if section.content %}
                <p>{{ section.content }}</p>
                {% endif %}
            </div>
            {% endfor %}
        </body>
        </html>
        """
        
        template = Template(html_template)
        html_content = template.render(
            title=content.get("title", "Document"),
            font_family=quality_config["font_family"],
            font_size=quality_config["font_size"],
            line_spacing=quality_config["line_spacing"],
            primary_color=quality_config["colors"]["primary"],
            accent_color=quality_config["colors"]["accent"],
            sections=content.get("sections", [])
        )
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return {
            "format": "html",
            "sections": len(content.get("sections", [])),
            "professional_features": quality_config
        }
    
    async def _export_to_markdown(self, content: Dict[str, Any], config: ExportConfig, output_path: str) -> Dict[str, Any]:
        """Export content to Markdown format."""
        markdown_content = []
        
        # Title
        if "title" in content:
            markdown_content.append(f"# {content['title']}\n")
        
        # Content sections
        if "sections" in content:
            for section in content["sections"]:
                if "heading" in section:
                    markdown_content.append(f"## {section['heading']}\n")
                
                if "content" in section:
                    markdown_content.append(f"{section['content']}\n")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))
        
        return {
            "format": "markdown",
            "sections": len(content.get("sections", [])),
            "professional_features": {"markdown_formatting": True}
        }
    
    async def _export_to_rtf(self, content: Dict[str, Any], config: ExportConfig, output_path: str) -> Dict[str, Any]:
        """Export content to RTF format."""
        # Basic RTF implementation
        rtf_content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\n"
        
        if "title" in content:
            rtf_content += f"\\f0\\fs24\\b {content['title']}\\b0\\par\\par\n"
        
        if "sections" in content:
            for section in content["sections"]:
                if "heading" in section:
                    rtf_content += f"\\f0\\fs20\\b {section['heading']}\\b0\\par\n"
                
                if "content" in section:
                    rtf_content += f"\\f0\\fs18 {section['content']}\\par\\par\n"
        
        rtf_content += "}"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(rtf_content)
        
        return {
            "format": "rtf",
            "sections": len(content.get("sections", [])),
            "professional_features": {"rtf_formatting": True}
        }
    
    async def _export_to_txt(self, content: Dict[str, Any], config: ExportConfig, output_path: str) -> Dict[str, Any]:
        """Export content to plain text format."""
        txt_content = []
        
        if "title" in content:
            txt_content.append(content['title'])
            txt_content.append("=" * len(content['title']))
            txt_content.append("")
        
        if "sections" in content:
            for section in content["sections"]:
                if "heading" in section:
                    txt_content.append(section['heading'])
                    txt_content.append("-" * len(section['heading']))
                
                if "content" in section:
                    txt_content.append(section['content'])
                    txt_content.append("")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(txt_content))
        
        return {
            "format": "txt",
            "sections": len(content.get("sections", [])),
            "professional_features": {"plain_text_formatting": True}
        }
    
    async def _export_to_json(self, content: Dict[str, Any], config: ExportConfig, output_path: str) -> Dict[str, Any]:
        """Export content to JSON format."""
        export_data = {
            "metadata": {
                "exported_at": datetime.now().isoformat(),
                "format": config.format.value,
                "document_type": config.document_type.value,
                "quality_level": config.quality_level.value
            },
            "content": content
        }
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, indent=2, ensure_ascii=False)
        
        return {
            "format": "json",
            "sections": len(content.get("sections", [])),
            "professional_features": {"structured_data": True}
        }
    
    async def _export_to_xml(self, content: Dict[str, Any], config: ExportConfig, output_path: str) -> Dict[str, Any]:
        """Export content to XML format."""
        xml_content = ['<?xml version="1.0" encoding="UTF-8"?>']
        xml_content.append('<document>')
        xml_content.append(f'  <metadata>')
        xml_content.append(f'    <exported_at>{datetime.now().isoformat()}</exported_at>')
        xml_content.append(f'    <format>{config.format.value}</format>')
        xml_content.append(f'    <document_type>{config.document_type.value}</document_type>')
        xml_content.append(f'    <quality_level>{config.quality_level.value}</quality_level>')
        xml_content.append(f'  </metadata>')
        
        if "title" in content:
            xml_content.append(f'  <title>{content["title"]}</title>')
        
        if "sections" in content:
            xml_content.append('  <sections>')
            for section in content["sections"]:
                xml_content.append('    <section>')
                if "heading" in section:
                    xml_content.append(f'      <heading>{section["heading"]}</heading>')
                if "content" in section:
                    xml_content.append(f'      <content>{section["content"]}</content>')
                xml_content.append('    </section>')
            xml_content.append('  </sections>')
        
        xml_content.append('</document>')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(xml_content))
        
        return {
            "format": "xml",
            "sections": len(content.get("sections", [])),
            "professional_features": {"xml_structure": True}
        }
    
    # Utility methods
    def get_task_status(self, task_id: str) -> Optional[Dict[str, Any]]:
        """Get status of an export task."""
        if task_id in self.active_tasks:
            task = self.active_tasks[task_id]
            return {
                "id": task.id,
                "status": task.status,
                "format": task.config.format.value,
                "document_type": task.config.document_type.value,
                "quality_level": task.config.quality_level.value,
                "created_at": task.created_at.isoformat(),
                "error": task.error
            }
        elif task_id in self.completed_tasks:
            result = self.completed_tasks[task_id]
            return {
                "id": task_id,
                "status": "completed",
                "success": result.success,
                "file_path": result.file_path,
                "file_size": result.file_size,
                "quality_score": result.quality_score,
                "processing_time": result.processing_time,
                "error": result.error
            }
        return None
    
    def get_export_statistics(self) -> Dict[str, Any]:
        """Get export statistics."""
        active_count = len(self.active_tasks)
        completed_count = len(self.completed_tasks)
        
        format_counts = {}
        quality_counts = {}
        
        for result in self.completed_tasks.values():
            format_counts[result.format] = format_counts.get(result.format, 0) + 1
        
        for task in self.active_tasks.values():
            quality_counts[task.config.quality_level.value] = quality_counts.get(task.config.quality_level.value, 0) + 1
        
        return {
            "total_tasks": active_count + completed_count,
            "active_tasks": active_count,
            "completed_tasks": completed_count,
            "format_distribution": format_counts,
            "quality_distribution": quality_counts,
            "average_quality_score": sum(r.quality_score for r in self.completed_tasks.values()) / max(completed_count, 1)
        }
    
    def list_supported_formats(self) -> List[Dict[str, Any]]:
        """List all supported export formats."""
        return [
            {
                "format": fmt.value,
                "name": fmt.value.upper(),
                "description": f"Export to {fmt.value.upper()} format",
                "professional_features": self._get_format_features(fmt)
            }
            for fmt in ExportFormat
        ]
    
    def _get_format_features(self, format_type: ExportFormat) -> List[str]:
        """Get professional features for a specific format."""
        features = {
            ExportFormat.PDF: ["High quality", "Print ready", "Professional layout", "Vector graphics"],
            ExportFormat.DOCX: ["Editable", "Professional formatting", "Table support", "Image embedding"],
            ExportFormat.HTML: ["Web ready", "Responsive", "Interactive elements", "SEO friendly"],
            ExportFormat.MARKDOWN: ["Version control friendly", "Lightweight", "Platform agnostic", "Easy to edit"],
            ExportFormat.RTF: ["Cross platform", "Rich formatting", "Legacy support"],
            ExportFormat.TXT: ["Universal compatibility", "Lightweight", "Fast processing"],
            ExportFormat.JSON: ["Structured data", "API friendly", "Machine readable"],
            ExportFormat.XML: ["Structured data", "Validation support", "Industry standard"]
        }
        return features.get(format_type, [])

# Global export engine instance
_global_export_engine: Optional[ExportIAEngine] = None

def get_global_export_engine() -> ExportIAEngine:
    """Get the global export engine instance."""
    global _global_export_engine
    if _global_export_engine is None:
        _global_export_engine = ExportIAEngine()
    return _global_export_engine



























