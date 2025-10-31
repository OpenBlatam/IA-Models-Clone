"""
Gamma App - Document Engine
Advanced AI-powered document generation with multiple formats
"""

import asyncio
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, asdict
from enum import Enum
import io
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import markdown
from markdown.extensions import tables, codehilite, fenced_code
import requests
from PIL import Image

logger = logging.getLogger(__name__)

class DocumentType(Enum):
    """Types of documents that can be generated"""
    REPORT = "report"
    PROPOSAL = "proposal"
    MANUAL = "manual"
    WHITEPAPER = "whitepaper"
    CASE_STUDY = "case_study"
    BUSINESS_PLAN = "business_plan"
    TECHNICAL_DOC = "technical_doc"
    USER_GUIDE = "user_guide"

class DocumentStyle(Enum):
    """Document styling options"""
    ACADEMIC = "academic"
    BUSINESS = "business"
    TECHNICAL = "technical"
    CREATIVE = "creative"
    FORMAL = "formal"
    CASUAL = "casual"

@dataclass
class DocumentSection:
    """Document section structure"""
    title: str
    content: str
    level: int  # Heading level (1-6)
    subsections: List['DocumentSection'] = None
    metadata: Dict[str, Any] = None

@dataclass
class DocumentTemplate:
    """Document template structure"""
    name: str
    type: DocumentType
    sections: List[DocumentSection]
    style: DocumentStyle
    metadata: Dict[str, Any]

class DocumentEngine:
    """
    Advanced document generation engine
    """
    
    def __init__(self, config: Optional[Dict] = None):
        """Initialize the document engine"""
        self.config = config or {}
        self.templates = {}
        self.styles = {}
        
        # Load templates and styles
        self._load_templates()
        self._load_styles()
        
        logger.info("Document Engine initialized successfully")

    def _load_templates(self):
        """Load document templates"""
        self.templates = {
            DocumentType.REPORT: {
                "sections": [
                    {"title": "Executive Summary", "level": 1},
                    {"title": "Introduction", "level": 1},
                    {"title": "Methodology", "level": 1},
                    {"title": "Findings", "level": 1},
                    {"title": "Analysis", "level": 1},
                    {"title": "Recommendations", "level": 1},
                    {"title": "Conclusion", "level": 1},
                    {"title": "References", "level": 1},
                    {"title": "Appendix", "level": 1}
                ],
                "style": DocumentStyle.BUSINESS
            },
            DocumentType.PROPOSAL: {
                "sections": [
                    {"title": "Project Overview", "level": 1},
                    {"title": "Problem Statement", "level": 1},
                    {"title": "Proposed Solution", "level": 1},
                    {"title": "Methodology", "level": 1},
                    {"title": "Timeline", "level": 1},
                    {"title": "Budget", "level": 1},
                    {"title": "Team", "level": 1},
                    {"title": "Expected Outcomes", "level": 1},
                    {"title": "Next Steps", "level": 1}
                ],
                "style": DocumentStyle.BUSINESS
            },
            DocumentType.MANUAL: {
                "sections": [
                    {"title": "Introduction", "level": 1},
                    {"title": "Getting Started", "level": 1},
                    {"title": "Basic Operations", "level": 1},
                    {"title": "Advanced Features", "level": 1},
                    {"title": "Troubleshooting", "level": 1},
                    {"title": "FAQ", "level": 1},
                    {"title": "Glossary", "level": 1},
                    {"title": "Index", "level": 1}
                ],
                "style": DocumentStyle.TECHNICAL
            },
            DocumentType.WHITEPAPER: {
                "sections": [
                    {"title": "Abstract", "level": 1},
                    {"title": "Introduction", "level": 1},
                    {"title": "Background", "level": 1},
                    {"title": "Problem Analysis", "level": 1},
                    {"title": "Solution Overview", "level": 1},
                    {"title": "Technical Details", "level": 1},
                    {"title": "Implementation", "level": 1},
                    {"title": "Results", "level": 1},
                    {"title": "Conclusion", "level": 1},
                    {"title": "References", "level": 1}
                ],
                "style": DocumentStyle.ACADEMIC
            }
        }

    def _load_styles(self):
        """Load document styles"""
        self.styles = {
            DocumentStyle.ACADEMIC: {
                "font_family": "Times New Roman",
                "font_size": 12,
                "line_spacing": 1.5,
                "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1},
                "header_style": "formal",
                "citation_style": "APA"
            },
            DocumentStyle.BUSINESS: {
                "font_family": "Arial",
                "font_size": 11,
                "line_spacing": 1.15,
                "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1},
                "header_style": "professional",
                "citation_style": "business"
            },
            DocumentStyle.TECHNICAL: {
                "font_family": "Calibri",
                "font_size": 11,
                "line_spacing": 1.2,
                "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1},
                "header_style": "technical",
                "citation_style": "IEEE"
            },
            DocumentStyle.CREATIVE: {
                "font_family": "Georgia",
                "font_size": 12,
                "line_spacing": 1.4,
                "margins": {"top": 1.2, "bottom": 1.2, "left": 1.2, "right": 1.2},
                "header_style": "creative",
                "citation_style": "creative"
            }
        }

    async def create_document(self, content: Dict[str, Any], 
                            doc_type: DocumentType = DocumentType.REPORT,
                            style: DocumentStyle = DocumentStyle.BUSINESS,
                            output_format: str = "docx") -> bytes:
        """Create a document from content"""
        try:
            if output_format == "docx":
                return await self._create_docx_document(content, doc_type, style)
            elif output_format == "pdf":
                return await self._create_pdf_document(content, doc_type, style)
            elif output_format == "html":
                return await self._create_html_document(content, doc_type, style)
            elif output_format == "markdown":
                return await self._create_markdown_document(content, doc_type, style)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")
                
        except Exception as e:
            logger.error(f"Error creating document: {e}")
            raise

    async def _create_docx_document(self, content: Dict[str, Any], 
                                  doc_type: DocumentType, style: DocumentStyle) -> bytes:
        """Create DOCX document"""
        try:
            # Create new document
            doc = Document()
            
            # Apply document style
            style_config = self.styles[style]
            self._apply_docx_style(doc, style_config)
            
            # Add title page
            await self._add_title_page_docx(doc, content, style_config)
            
            # Add table of contents
            await self._add_table_of_contents_docx(doc, content, doc_type)
            
            # Add main content
            await self._add_main_content_docx(doc, content, doc_type, style_config)
            
            # Add references if needed
            if doc_type in [DocumentType.REPORT, DocumentType.WHITEPAPER]:
                await self._add_references_docx(doc, content, style_config)
            
            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            logger.info("DOCX document created successfully")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error creating DOCX document: {e}")
            raise

    async def _create_pdf_document(self, content: Dict[str, Any], 
                                 doc_type: DocumentType, style: DocumentStyle) -> bytes:
        """Create PDF document"""
        try:
            output = io.BytesIO()
            doc = SimpleDocTemplate(output, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Apply custom styles
            style_config = self.styles[style]
            self._apply_pdf_styles(styles, style_config)
            
            # Add title page
            await self._add_title_page_pdf(story, content, styles)
            
            # Add table of contents
            await self._add_table_of_contents_pdf(story, content, doc_type, styles)
            
            # Add main content
            await self._add_main_content_pdf(story, content, doc_type, styles)
            
            # Add references if needed
            if doc_type in [DocumentType.REPORT, DocumentType.WHITEPAPER]:
                await self._add_references_pdf(story, content, styles)
            
            doc.build(story)
            output.seek(0)
            
            logger.info("PDF document created successfully")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error creating PDF document: {e}")
            raise

    async def _create_html_document(self, content: Dict[str, Any], 
                                  doc_type: DocumentType, style: DocumentStyle) -> bytes:
        """Create HTML document"""
        try:
            html_content = f"""
            <!DOCTYPE html>
            <html lang="en">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{content.get('title', 'Document')}</title>
                <style>
                    {self._get_html_css(style)}
                </style>
            </head>
            <body>
                <div class="document">
                    {await self._generate_html_content(content, doc_type)}
                </div>
            </body>
            </html>
            """
            
            logger.info("HTML document created successfully")
            return html_content.encode('utf-8')
            
        except Exception as e:
            logger.error(f"Error creating HTML document: {e}")
            raise

    async def _create_markdown_document(self, content: Dict[str, Any], 
                                      doc_type: DocumentType, style: DocumentStyle) -> bytes:
        """Create Markdown document"""
        try:
            markdown_content = f"# {content.get('title', 'Document')}\n\n"
            
            if 'subtitle' in content:
                markdown_content += f"*{content['subtitle']}*\n\n"
            
            # Add main content
            if 'sections' in content:
                for section in content['sections']:
                    markdown_content += f"## {section.get('section_name', 'Section').title()}\n\n"
                    markdown_content += f"{section.get('content', '')}\n\n"
            
            logger.info("Markdown document created successfully")
            return markdown_content.encode('utf-8')
            
        except Exception as e:
            logger.error(f"Error creating Markdown document: {e}")
            raise

    def _apply_docx_style(self, doc: Document, style_config: Dict[str, Any]):
        """Apply style to DOCX document"""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = style_config['font_family']
        font.size = Pt(style_config['font_size'])
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(style_config['margins']['top'])
            section.bottom_margin = Inches(style_config['margins']['bottom'])
            section.left_margin = Inches(style_config['margins']['left'])
            section.right_margin = Inches(style_config['margins']['right'])

    def _apply_pdf_styles(self, styles, style_config: Dict[str, Any]):
        """Apply custom styles to PDF"""
        # Create custom styles based on document style
        custom_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=style_config['font_family'],
            fontSize=style_config['font_size'],
            leading=style_config['font_size'] * style_config['line_spacing']
        )
        styles.add(custom_style)

    async def _add_title_page_docx(self, doc: Document, content: Dict[str, Any], 
                                 style_config: Dict[str, Any]):
        """Add title page to DOCX document"""
        # Add title
        title = doc.add_heading(content.get('title', 'Document Title'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle if available
        if 'subtitle' in content:
            subtitle = doc.add_paragraph(content['subtitle'])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add author and date
        doc.add_paragraph(f"Author: {content.get('author', 'Generated by Gamma App')}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        
        # Add page break
        doc.add_page_break()

    async def _add_title_page_pdf(self, story: List, content: Dict[str, Any], styles):
        """Add title page to PDF document"""
        # Add title
        title_style = styles['Title']
        title_style.alignment = 1  # Center alignment
        story.append(Paragraph(content.get('title', 'Document Title'), title_style))
        story.append(Spacer(1, 12))
        
        # Add subtitle if available
        if 'subtitle' in content:
            subtitle_style = styles['Heading2']
            subtitle_style.alignment = 1
            story.append(Paragraph(content['subtitle'], subtitle_style))
            story.append(Spacer(1, 12))
        
        # Add author and date
        story.append(Paragraph(f"Author: {content.get('author', 'Generated by Gamma App')}", styles['Normal']))
        story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(PageBreak())

    async def _add_table_of_contents_docx(self, doc: Document, content: Dict[str, Any], 
                                        doc_type: DocumentType):
        """Add table of contents to DOCX document"""
        toc_heading = doc.add_heading('Table of Contents', level=1)
        
        # Get template sections
        template = self.templates.get(doc_type, self.templates[DocumentType.REPORT])
        
        for section in template['sections']:
            doc.add_paragraph(f"{section['title']} ................. {len(doc.paragraphs) + 1}")

    async def _add_table_of_contents_pdf(self, story: List, content: Dict[str, Any], 
                                       doc_type: DocumentType, styles):
        """Add table of contents to PDF document"""
        story.append(Paragraph('Table of Contents', styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Get template sections
        template = self.templates.get(doc_type, self.templates[DocumentType.REPORT])
        
        for section in template['sections']:
            story.append(Paragraph(f"{section['title']} ................. {len(story) + 1}", styles['Normal']))
        
        story.append(PageBreak())

    async def _add_main_content_docx(self, doc: Document, content: Dict[str, Any], 
                                   doc_type: DocumentType, style_config: Dict[str, Any]):
        """Add main content to DOCX document"""
        template = self.templates.get(doc_type, self.templates[DocumentType.REPORT])
        
        for section in template['sections']:
            # Add section heading
            heading = doc.add_heading(section['title'], level=section['level'])
            
            # Add section content
            section_content = self._get_section_content(content, section['title'])
            if section_content:
                doc.add_paragraph(section_content)
            else:
                doc.add_paragraph(f"Content for {section['title']} will be generated here.")

    async def _add_main_content_pdf(self, story: List, content: Dict[str, Any], 
                                  doc_type: DocumentType, styles):
        """Add main content to PDF document"""
        template = self.templates.get(doc_type, self.templates[DocumentType.REPORT])
        
        for section in template['sections']:
            # Add section heading
            heading_style = styles[f'Heading{section["level"]}']
            story.append(Paragraph(section['title'], heading_style))
            story.append(Spacer(1, 12))
            
            # Add section content
            section_content = self._get_section_content(content, section['title'])
            if section_content:
                story.append(Paragraph(section_content, styles['Normal']))
            else:
                story.append(Paragraph(f"Content for {section['title']} will be generated here.", styles['Normal']))
            
            story.append(Spacer(1, 12))

    async def _add_references_docx(self, doc: Document, content: Dict[str, Any], 
                                 style_config: Dict[str, Any]):
        """Add references section to DOCX document"""
        ref_heading = doc.add_heading('References', level=1)
        
        # Add sample references
        references = [
            "Smith, J. (2023). Advanced AI Applications. Journal of Technology, 15(3), 45-62.",
            "Johnson, A. (2023). Machine Learning in Business. Business Review, 28(2), 12-25.",
            "Brown, M. (2023). Future of Artificial Intelligence. Tech Today, 10(1), 8-15."
        ]
        
        for ref in references:
            doc.add_paragraph(ref, style='List Number')

    async def _add_references_pdf(self, story: List, content: Dict[str, Any], styles):
        """Add references section to PDF document"""
        story.append(Paragraph('References', styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Add sample references
        references = [
            "Smith, J. (2023). Advanced AI Applications. Journal of Technology, 15(3), 45-62.",
            "Johnson, A. (2023). Machine Learning in Business. Business Review, 28(2), 12-25.",
            "Brown, M. (2023). Future of Artificial Intelligence. Tech Today, 10(1), 8-15."
        ]
        
        for ref in references:
            story.append(Paragraph(ref, styles['Normal']))

    def _get_section_content(self, content: Dict[str, Any], section_title: str) -> str:
        """Get content for a specific section"""
        if 'sections' in content:
            for section in content['sections']:
                if section.get('section_name', '').lower() == section_title.lower().replace(' ', '_'):
                    return section.get('content', '')
        
        return ""

    def _get_html_css(self, style: DocumentStyle) -> str:
        """Get CSS styles for HTML document"""
        style_config = self.styles[style]
        
        return f"""
        body {{
            font-family: {style_config['font_family']};
            font-size: {style_config['font_size']}pt;
            line-height: {style_config['line_spacing']};
            margin: {style_config['margins']['top']}in {style_config['margins']['right']}in {style_config['margins']['bottom']}in {style_config['margins']['left']}in;
            color: #333;
        }}
        
        .document {{
            max-width: 8.5in;
            margin: 0 auto;
        }}
        
        h1, h2, h3, h4, h5, h6 {{
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }}
        
        h1 {{ font-size: 24pt; }}
        h2 {{ font-size: 20pt; }}
        h3 {{ font-size: 16pt; }}
        
        p {{
            margin-bottom: 1em;
            text-align: justify;
        }}
        
        .toc {{
            margin: 2em 0;
        }}
        
        .toc ul {{
            list-style: none;
            padding-left: 0;
        }}
        
        .toc li {{
            margin: 0.5em 0;
        }}
        """

    async def _generate_html_content(self, content: Dict[str, Any], doc_type: DocumentType) -> str:
        """Generate HTML content"""
        html = f"<h1>{content.get('title', 'Document')}</h1>"
        
        if 'subtitle' in content:
            html += f"<p class='subtitle'>{content['subtitle']}</p>"
        
        # Add table of contents
        html += "<div class='toc'><h2>Table of Contents</h2><ul>"
        template = self.templates.get(doc_type, self.templates[DocumentType.REPORT])
        for section in template['sections']:
            html += f"<li><a href='#{section['title'].replace(' ', '_')}'>{section['title']}</a></li>"
        html += "</ul></div>"
        
        # Add main content
        if 'sections' in content:
            for section in content['sections']:
                section_id = section.get('section_name', 'section').replace(' ', '_')
                html += f"<h2 id='{section_id}'>{section.get('section_name', 'Section').title()}</h2>"
                html += f"<p>{section.get('content', '')}</p>"
        
        return html

    def get_available_document_types(self) -> List[DocumentType]:
        """Get available document types"""
        return list(DocumentType)

    def get_available_styles(self) -> List[DocumentStyle]:
        """Get available document styles"""
        return list(DocumentStyle)

    def get_available_output_formats(self) -> List[str]:
        """Get available output formats"""
        return ["docx", "pdf", "html", "markdown"]

    async def add_table_to_document(self, document_bytes: bytes, 
                                  table_data: List[List[str]], 
                                  output_format: str = "docx") -> bytes:
        """Add table to existing document"""
        # This would implement table addition to documents
        # Implementation would depend on specific requirements
        return document_bytes

    async def add_image_to_document(self, document_bytes: bytes, 
                                  image_data: bytes, 
                                  caption: str = "",
                                  output_format: str = "docx") -> bytes:
        """Add image to existing document"""
        # This would implement image addition to documents
        # Implementation would depend on specific requirements
        return document_bytes

    async def add_chart_to_document(self, document_bytes: bytes, 
                                  chart_data: Dict[str, Any], 
                                  chart_type: str = "bar",
                                  output_format: str = "docx") -> bytes:
        """Add chart to existing document"""
        # This would implement chart addition to documents
        # Implementation would depend on specific requirements
        return document_bytes



























