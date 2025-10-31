"""
Professional Document Services
==============================

Core services for document generation, processing, and export functionality.
"""

import asyncio
import logging
import os
import tempfile
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from uuid import uuid4

import aiofiles
from jinja2 import Environment, FileSystemLoader, Template
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.platypus.tableofcontents import TableOfContents
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import (
    DocumentGenerationRequest,
    DocumentGenerationResponse,
    DocumentExportRequest,
    DocumentExportResponse,
    DocumentTemplate,
    ProfessionalDocument,
    DocumentSection,
    DocumentStyle,
    ExportFormat,
    DocumentType
)
from .templates import template_manager
from .ai_service import AIDocumentGenerator

logger = logging.getLogger(__name__)


class DocumentGenerationService:
    """Service for generating professional documents using AI."""
    
    def __init__(self):
        self.ai_generator = AIDocumentGenerator()
        self.documents: Dict[str, ProfessionalDocument] = {}
    
    async def generate_document(self, request: DocumentGenerationRequest) -> DocumentGenerationResponse:
        """Generate a professional document based on the request."""
        start_time = time.time()
        
        try:
            # Get template
            if request.template_id:
                template = template_manager.get_template(request.template_id)
            else:
                template = template_manager.get_default_template(request.document_type)
            
            # Create document structure
            document = ProfessionalDocument(
                title=request.title or self._generate_title(request.query, request.document_type),
                subtitle=request.subtitle,
                document_type=request.document_type,
                template_id=template.id,
                author=request.author,
                company=request.company,
                style=request.style or template.style,
                status="generating"
            )
            
            # Generate content using AI
            sections = await self.ai_generator.generate_document_content(
                query=request.query,
                template=template,
                document_type=request.document_type,
                tone=request.tone,
                length=request.length,
                language=request.language,
                additional_requirements=request.additional_requirements
            )
            
            # Process and structure sections
            document.sections = self._process_sections(sections, template)
            document.word_count = sum(len(section.content.split()) for section in document.sections)
            document.page_count = self._estimate_page_count(document.word_count, document.style)
            document.status = "completed"
            document.date_modified = datetime.utcnow()
            
            # Store document
            self.documents[document.id] = document
            
            generation_time = time.time() - start_time
            
            return DocumentGenerationResponse(
                success=True,
                document=document,
                message="Document generated successfully",
                generation_time=generation_time,
                word_count=document.word_count,
                estimated_pages=document.page_count
            )
            
        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            generation_time = time.time() - start_time
            
            return DocumentGenerationResponse(
                success=False,
                document=None,
                message=f"Error generating document: {str(e)}",
                generation_time=generation_time
            )
    
    def _generate_title(self, query: str, document_type: DocumentType) -> str:
        """Generate a title based on the query and document type."""
        # Simple title generation - in production, this could use AI
        query_words = query.split()[:5]  # Take first 5 words
        base_title = " ".join(query_words).title()
        
        type_suffixes = {
            DocumentType.REPORT: "Report",
            DocumentType.PROPOSAL: "Proposal",
            DocumentType.MANUAL: "Manual",
            DocumentType.GUIDE: "Guide",
            DocumentType.WHITEPAPER: "Whitepaper",
            DocumentType.BUSINESS_PLAN: "Business Plan",
            DocumentType.TECHNICAL_DOCUMENT: "Technical Documentation",
            DocumentType.ACADEMIC_PAPER: "Research Paper",
            DocumentType.NEWSLETTER: "Newsletter",
            DocumentType.BROCHURE: "Brochure",
            DocumentType.CATALOG: "Catalog",
            DocumentType.PRESENTATION: "Presentation"
        }
        
        suffix = type_suffixes.get(document_type, "Document")
        return f"{base_title} - {suffix}"
    
    def _process_sections(self, sections: List[Dict[str, Any]], template: DocumentTemplate) -> List[DocumentSection]:
        """Process AI-generated sections into DocumentSection objects."""
        processed_sections = []
        
        for i, section_data in enumerate(sections):
            section = DocumentSection(
                title=section_data.get("title", f"Section {i+1}"),
                content=section_data.get("content", ""),
                level=section_data.get("level", 1),
                order=i,
                metadata=section_data.get("metadata", {})
            )
            processed_sections.append(section)
        
        return processed_sections
    
    def _estimate_page_count(self, word_count: int, style: DocumentStyle) -> int:
        """Estimate page count based on word count and styling."""
        # Rough estimation: 250-300 words per page for standard formatting
        words_per_page = 275
        return max(1, word_count // words_per_page)
    
    def get_document(self, document_id: str) -> Optional[ProfessionalDocument]:
        """Get a document by ID."""
        return self.documents.get(document_id)
    
    def list_documents(self, limit: int = 50, offset: int = 0) -> List[ProfessionalDocument]:
        """List documents with pagination."""
        documents = list(self.documents.values())
        documents.sort(key=lambda x: x.date_created, reverse=True)
        return documents[offset:offset + limit]
    
    def update_document(self, document_id: str, updates: Dict[str, Any]) -> Optional[ProfessionalDocument]:
        """Update a document."""
        if document_id not in self.documents:
            return None
        
        document = self.documents[document_id]
        
        # Update fields
        for key, value in updates.items():
            if hasattr(document, key):
                setattr(document, key, value)
        
        document.date_modified = datetime.utcnow()
        return document


class DocumentExportService:
    """Service for exporting documents in various formats."""
    
    def __init__(self, output_dir: str = "exports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    async def export_document(
        self, 
        document: ProfessionalDocument, 
        request: DocumentExportRequest
    ) -> DocumentExportResponse:
        """Export a document in the specified format."""
        start_time = time.time()
        
        try:
            if request.format == ExportFormat.PDF:
                file_path = await self._export_to_pdf(document, request)
            elif request.format == ExportFormat.MARKDOWN:
                file_path = await self._export_to_markdown(document, request)
            elif request.format == ExportFormat.WORD:
                file_path = await self._export_to_word(document, request)
            elif request.format == ExportFormat.HTML:
                file_path = await self._export_to_html(document, request)
            else:
                raise ValueError(f"Unsupported export format: {request.format}")
            
            # Get file size
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            
            export_time = time.time() - start_time
            
            return DocumentExportResponse(
                success=True,
                file_path=str(file_path),
                file_size=file_size,
                download_url=f"/download/{Path(file_path).name}",
                message="Document exported successfully",
                export_time=export_time
            )
            
        except Exception as e:
            logger.error(f"Error exporting document: {str(e)}")
            export_time = time.time() - start_time
            
            return DocumentExportResponse(
                success=False,
                message=f"Error exporting document: {str(e)}",
                export_time=export_time
            )
    
    async def _export_to_pdf(self, document: ProfessionalDocument, request: DocumentExportRequest) -> str:
        """Export document to PDF format."""
        filename = request.custom_filename or f"{document.title.replace(' ', '_')}.pdf"
        file_path = self.output_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=A4,
            rightMargin=document.style.margin_right * inch,
            leftMargin=document.style.margin_left * inch,
            topMargin=document.style.margin_top * inch,
            bottomMargin=document.style.margin_bottom * inch
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles based on document style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            textColor=colors.HexColor(document.style.header_color),
            fontName=document.style.font_family
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.HexColor(document.style.header_color),
            fontName=document.style.font_family
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=document.style.font_size,
            spaceAfter=6,
            textColor=colors.HexColor(document.style.body_color),
            fontName=document.style.font_family,
            leading=document.style.font_size * document.style.line_spacing
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph(document.title, title_style))
        story.append(Spacer(1, 12))
        
        # Subtitle
        if document.subtitle:
            story.append(Paragraph(document.subtitle, heading_style))
            story.append(Spacer(1, 12))
        
        # Author and company
        if document.author or document.company:
            author_info = []
            if document.author:
                author_info.append(f"Author: {document.author}")
            if document.company:
                author_info.append(f"Company: {document.company}")
            story.append(Paragraph("<br/>".join(author_info), body_style))
            story.append(Spacer(1, 12))
        
        # Date
        story.append(Paragraph(f"Date: {document.date_created.strftime('%B %d, %Y')}", body_style))
        story.append(Spacer(1, 20))
        
        # Sections
        for section in document.sections:
            if section.level == 1:
                story.append(Paragraph(section.title, title_style))
            elif section.level == 2:
                story.append(Paragraph(section.title, heading_style))
            else:
                story.append(Paragraph(section.title, body_style))
            
            story.append(Spacer(1, 6))
            story.append(Paragraph(section.content, body_style))
            story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        
        return str(file_path)
    
    async def _export_to_markdown(self, document: ProfessionalDocument, request: DocumentExportRequest) -> str:
        """Export document to Markdown format."""
        filename = request.custom_filename or f"{document.title.replace(' ', '_')}.md"
        file_path = self.output_dir / filename
        
        markdown_content = []
        
        # Title
        markdown_content.append(f"# {document.title}")
        markdown_content.append("")
        
        # Subtitle
        if document.subtitle:
            markdown_content.append(f"## {document.subtitle}")
            markdown_content.append("")
        
        # Metadata
        if document.author or document.company:
            markdown_content.append("---")
            if document.author:
                markdown_content.append(f"**Author:** {document.author}")
            if document.company:
                markdown_content.append(f"**Company:** {document.company}")
            markdown_content.append(f"**Date:** {document.date_created.strftime('%B %d, %Y')}")
            markdown_content.append("---")
            markdown_content.append("")
        
        # Sections
        for section in document.sections:
            heading_level = "#" * min(section.level + 1, 6)
            markdown_content.append(f"{heading_level} {section.title}")
            markdown_content.append("")
            markdown_content.append(section.content)
            markdown_content.append("")
        
        # Write to file
        async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
            await f.write('\n'.join(markdown_content))
        
        return str(file_path)
    
    async def _export_to_word(self, document: ProfessionalDocument, request: DocumentExportRequest) -> str:
        """Export document to Word format."""
        filename = request.custom_filename or f"{document.title.replace(' ', '_')}.docx"
        file_path = self.output_dir / filename
        
        # Create Word document
        doc = DocxDocument()
        
        # Set document properties
        doc.core_properties.title = document.title
        if document.author:
            doc.core_properties.author = document.author
        doc.core_properties.created = document.date_created
        
        # Title
        title = doc.add_heading(document.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        if document.subtitle:
            subtitle = doc.add_heading(document.subtitle, 1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Author and company info
        if document.author or document.company:
            info_para = doc.add_paragraph()
            if document.author:
                info_para.add_run(f"Author: {document.author}\n")
            if document.company:
                info_para.add_run(f"Company: {document.company}\n")
            info_para.add_run(f"Date: {document.date_created.strftime('%B %d, %Y')}")
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        doc.add_page_break()
        
        # Sections
        for section in document.sections:
            # Add heading
            heading = doc.add_heading(section.title, section.level)
            
            # Add content
            content_para = doc.add_paragraph(section.content)
            
            # Apply formatting based on document style
            for run in content_para.runs:
                run.font.name = document.style.font_family
                run.font.size = Pt(document.style.font_size)
        
        # Save document
        doc.save(str(file_path))
        
        return str(file_path)
    
    async def _export_to_html(self, document: ProfessionalDocument, request: DocumentExportRequest) -> str:
        """Export document to HTML format."""
        filename = request.custom_filename or f"{document.title.replace(' ', '_')}.html"
        file_path = self.output_dir / filename
        
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{document.title}</title>
    <style>
        body {{
            font-family: {document.style.font_family};
            font-size: {document.style.font_size}px;
            line-height: {document.style.line_spacing};
            color: {document.style.body_color};
            background-color: {document.style.background_color};
            margin: {document.style.margin_top}in {document.style.margin_right}in {document.style.margin_bottom}in {document.style.margin_left}in;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: {document.style.header_color};
        }}
        .title {{
            text-align: center;
            margin-bottom: 2em;
        }}
        .metadata {{
            text-align: center;
            margin-bottom: 2em;
            font-style: italic;
        }}
        .section {{
            margin-bottom: 1.5em;
        }}
    </style>
</head>
<body>
    <div class="title">
        <h1>{document.title}</h1>
        {f'<h2>{document.subtitle}</h2>' if document.subtitle else ''}
    </div>
    
    <div class="metadata">
        {f'<p><strong>Author:</strong> {document.author}</p>' if document.author else ''}
        {f'<p><strong>Company:</strong> {document.company}</p>' if document.company else ''}
        <p><strong>Date:</strong> {document.date_created.strftime('%B %d, %Y')}</p>
    </div>
    
    <div class="content">
"""
        
        # Add sections
        for section in document.sections:
            heading_tag = f"h{min(section.level + 1, 6)}"
            html_content += f"""
        <div class="section">
            <{heading_tag}>{section.title}</{heading_tag}>
            <p>{section.content.replace(chr(10), '<br>')}</p>
        </div>
"""
        
        html_content += """
    </div>
</body>
</html>
"""
        
        # Write to file
        async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
            await f.write(html_content)
        
        return str(file_path)


class TemplateService:
    """Service for managing document templates."""
    
    def __init__(self):
        self.template_manager = template_manager
    
    def get_template(self, template_id: str) -> DocumentTemplate:
        """Get a template by ID."""
        return self.template_manager.get_template(template_id)
    
    def get_templates_by_type(self, document_type: DocumentType) -> List[DocumentTemplate]:
        """Get templates by document type."""
        return self.template_manager.get_templates_by_type(document_type)
    
    def get_all_templates(self) -> List[DocumentTemplate]:
        """Get all available templates."""
        return self.template_manager.get_all_templates()
    
    def get_default_template(self, document_type: DocumentType) -> DocumentTemplate:
        """Get default template for document type."""
        return self.template_manager.get_default_template(document_type)
    
    def add_custom_template(self, template: DocumentTemplate):
        """Add a custom template."""
        self.template_manager.add_custom_template(template)
    
    def remove_template(self, template_id: str):
        """Remove a template."""
        self.template_manager.remove_template(template_id)




























