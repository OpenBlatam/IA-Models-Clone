"""
Advanced Export Service for Professional Documents
=================================================

Enhanced export service with advanced features like watermarks, digital signatures,
interactive elements, and custom branding.
"""

import asyncio
import os
import tempfile
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from uuid import uuid4

import aiofiles
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.graphics.shapes import Drawing, Rect
from reportlab.graphics import renderPDF
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from .models import (
    ProfessionalDocument,
    DocumentExportRequest,
    DocumentExportResponse,
    ExportFormat,
    DocumentStyle
)
from .services import DocumentExportService

logger = logging.getLogger(__name__)


class WatermarkType(str, Enum):
    """Types of watermarks."""
    TEXT = "text"
    IMAGE = "image"
    LOGO = "logo"
    CONFIDENTIAL = "confidential"
    DRAFT = "draft"
    CUSTOM = "custom"


class SignatureType(str, Enum):
    """Types of digital signatures."""
    TEXT = "text"
    IMAGE = "image"
    DIGITAL = "digital"
    ELECTRONIC = "electronic"


class AdvancedExportService(DocumentExportService):
    """Advanced export service with enhanced features."""
    
    def __init__(self, output_dir: str = "exports"):
        super().__init__(output_dir)
        self.watermark_cache = {}
        self.signature_cache = {}
        self.branding_cache = {}
    
    async def export_document_advanced(
        self,
        document: ProfessionalDocument,
        request: DocumentExportRequest,
        watermark_config: Optional[Dict[str, Any]] = None,
        signature_config: Optional[Dict[str, Any]] = None,
        branding_config: Optional[Dict[str, Any]] = None,
        interactive_features: Optional[Dict[str, Any]] = None
    ) -> DocumentExportResponse:
        """Export document with advanced features."""
        
        start_time = time.time()
        
        try:
            # Apply advanced configurations
            if watermark_config:
                document = await self._apply_watermark_config(document, watermark_config)
            
            if signature_config:
                document = await self._apply_signature_config(document, signature_config)
            
            if branding_config:
                document = await self._apply_branding_config(document, branding_config)
            
            if interactive_features:
                document = await self._apply_interactive_features(document, interactive_features)
            
            # Export with enhanced features
            if request.format == ExportFormat.PDF:
                file_path = await self._export_to_advanced_pdf(document, request, watermark_config, signature_config, branding_config)
            elif request.format == ExportFormat.WORD:
                file_path = await self._export_to_advanced_word(document, request, watermark_config, signature_config, branding_config)
            elif request.format == ExportFormat.HTML:
                file_path = await self._export_to_advanced_html(document, request, interactive_features)
            else:
                # Fall back to standard export
                file_path = await self._export_document_standard(document, request)
            
            # Get file size
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            
            export_time = time.time() - start_time
            
            return DocumentExportResponse(
                success=True,
                file_path=str(file_path),
                file_size=file_size,
                download_url=f"/download/{Path(file_path).name}",
                message="Document exported successfully with advanced features",
                export_time=export_time
            )
            
        except Exception as e:
            logger.error(f"Error in advanced export: {str(e)}")
            export_time = time.time() - start_time
            
            return DocumentExportResponse(
                success=False,
                message=f"Advanced export failed: {str(e)}",
                export_time=export_time
            )
    
    async def _apply_watermark_config(self, document: ProfessionalDocument, config: Dict[str, Any]) -> ProfessionalDocument:
        """Apply watermark configuration to document."""
        
        watermark_type = config.get("type", WatermarkType.TEXT)
        watermark_text = config.get("text", "CONFIDENTIAL")
        watermark_opacity = config.get("opacity", 0.3)
        watermark_position = config.get("position", "center")
        watermark_rotation = config.get("rotation", 45)
        
        # Store watermark config in document metadata
        if not document.metadata:
            document.metadata = {}
        
        document.metadata["watermark"] = {
            "type": watermark_type,
            "text": watermark_text,
            "opacity": watermark_opacity,
            "position": watermark_position,
            "rotation": watermark_rotation
        }
        
        return document
    
    async def _apply_signature_config(self, document: ProfessionalDocument, config: Dict[str, Any]) -> ProfessionalDocument:
        """Apply signature configuration to document."""
        
        signature_type = config.get("type", SignatureType.TEXT)
        signature_text = config.get("text", "Digitally Signed")
        signature_image_path = config.get("image_path")
        signature_position = config.get("position", "bottom_right")
        signature_date = config.get("include_date", True)
        
        # Store signature config in document metadata
        if not document.metadata:
            document.metadata = {}
        
        document.metadata["signature"] = {
            "type": signature_type,
            "text": signature_text,
            "image_path": signature_image_path,
            "position": signature_position,
            "include_date": signature_date,
            "signature_date": datetime.now().isoformat() if signature_date else None
        }
        
        return document
    
    async def _apply_branding_config(self, document: ProfessionalDocument, config: Dict[str, Any]) -> ProfessionalDocument:
        """Apply branding configuration to document."""
        
        logo_path = config.get("logo_path")
        company_name = config.get("company_name", document.company)
        company_address = config.get("company_address")
        company_phone = config.get("company_phone")
        company_email = config.get("company_email")
        website = config.get("website")
        color_scheme = config.get("color_scheme", {})
        
        # Store branding config in document metadata
        if not document.metadata:
            document.metadata = {}
        
        document.metadata["branding"] = {
            "logo_path": logo_path,
            "company_name": company_name,
            "company_address": company_address,
            "company_phone": company_phone,
            "company_email": company_email,
            "website": website,
            "color_scheme": color_scheme
        }
        
        return document
    
    async def _apply_interactive_features(self, document: ProfessionalDocument, features: Dict[str, Any]) -> ProfessionalDocument:
        """Apply interactive features to document."""
        
        enable_hyperlinks = features.get("hyperlinks", True)
        enable_bookmarks = features.get("bookmarks", True)
        enable_forms = features.get("forms", False)
        enable_annotations = features.get("annotations", False)
        enable_search = features.get("search", True)
        
        # Store interactive features in document metadata
        if not document.metadata:
            document.metadata = {}
        
        document.metadata["interactive_features"] = {
            "hyperlinks": enable_hyperlinks,
            "bookmarks": enable_bookmarks,
            "forms": enable_forms,
            "annotations": enable_annotations,
            "search": enable_search
        }
        
        return document
    
    async def _export_to_advanced_pdf(
        self,
        document: ProfessionalDocument,
        request: DocumentExportRequest,
        watermark_config: Optional[Dict[str, Any]],
        signature_config: Optional[Dict[str, Any]],
        branding_config: Optional[Dict[str, Any]]
    ) -> str:
        """Export document to PDF with advanced features."""
        
        filename = request.custom_filename or f"{document.title.replace(' ', '_')}_advanced.pdf"
        file_path = self.output_dir / filename
        
        # Create PDF document
        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=A4,
            rightMargin=document.style.margin_right * inch,
            leftMargin=document.style.margin_left * inch,
            topMargin=document.style.margin_top * inch,
            bottomMargin=document.style.margin_bottom * inch
        )
        
        # Define enhanced styles
        styles = getSampleStyleSheet()
        
        # Enhanced title style
        title_style = ParagraphStyle(
            'EnhancedTitle',
            parent=styles['Heading1'],
            fontSize=20,
            spaceAfter=30,
            textColor=colors.HexColor(document.style.header_color),
            fontName=document.style.font_family,
            alignment=TA_CENTER
        )
        
        # Enhanced heading style
        heading_style = ParagraphStyle(
            'EnhancedHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.HexColor(document.style.header_color),
            fontName=document.style.font_family,
            borderWidth=1,
            borderColor=colors.HexColor(document.style.accent_color),
            borderPadding=5
        )
        
        # Enhanced body style
        body_style = ParagraphStyle(
            'EnhancedBody',
            parent=styles['Normal'],
            fontSize=document.style.font_size,
            spaceAfter=6,
            textColor=colors.HexColor(document.style.body_color),
            fontName=document.style.font_family,
            leading=document.style.font_size * document.style.line_spacing
        )
        
        # Build content
        story = []
        
        # Add branding header if configured
        if branding_config and document.metadata.get("branding"):
            story.extend(await self._create_branding_header(document))
        
        # Title
        story.append(Paragraph(document.title, title_style))
        story.append(Spacer(1, 12))
        
        # Subtitle
        if document.subtitle:
            story.append(Paragraph(document.subtitle, heading_style))
            story.append(Spacer(1, 12))
        
        # Author and company info
        if document.author or document.company:
            author_info = []
            if document.author:
                author_info.append(f"<b>Author:</b> {document.author}")
            if document.company:
                author_info.append(f"<b>Company:</b> {document.company}")
            story.append(Paragraph("<br/>".join(author_info), body_style))
            story.append(Spacer(1, 12))
        
        # Date
        story.append(Paragraph(f"<b>Date:</b> {document.date_created.strftime('%B %d, %Y')}", body_style))
        story.append(Spacer(1, 20))
        
        # Table of Contents if document is long
        if len(document.sections) > 5:
            story.extend(await self._create_table_of_contents(document))
        
        # Sections with enhanced formatting
        for section in document.sections:
            if section.level == 1:
                story.append(Paragraph(section.title, title_style))
            elif section.level == 2:
                story.append(Paragraph(section.title, heading_style))
            else:
                story.append(Paragraph(section.title, body_style))
            
            story.append(Spacer(1, 6))
            
            # Enhanced content with formatting
            enhanced_content = await self._enhance_content_formatting(section.content)
            story.append(Paragraph(enhanced_content, body_style))
            story.append(Spacer(1, 12))
        
        # Add signature if configured
        if signature_config and document.metadata.get("signature"):
            story.extend(await self._create_signature_section(document))
        
        # Build PDF with watermark
        if watermark_config and document.metadata.get("watermark"):
            doc.build(story, onFirstPage=self._add_watermark, onLaterPages=self._add_watermark)
        else:
            doc.build(story)
        
        return str(file_path)
    
    async def _create_branding_header(self, document: ProfessionalDocument) -> List:
        """Create branding header for document."""
        story = []
        
        branding = document.metadata.get("branding", {})
        
        if branding.get("logo_path") and os.path.exists(branding["logo_path"]):
            # Add logo (simplified - would need proper image handling)
            story.append(Paragraph(f"<b>{branding.get('company_name', '')}</b>", getSampleStyleSheet()['Heading3']))
        
        if branding.get("company_address"):
            story.append(Paragraph(branding["company_address"], getSampleStyleSheet()['Normal']))
        
        if branding.get("company_phone") or branding.get("company_email"):
            contact_info = []
            if branding.get("company_phone"):
                contact_info.append(f"Phone: {branding['company_phone']}")
            if branding.get("company_email"):
                contact_info.append(f"Email: {branding['company_email']}")
            story.append(Paragraph(" | ".join(contact_info), getSampleStyleSheet()['Normal']))
        
        story.append(Spacer(1, 20))
        
        return story
    
    async def _create_table_of_contents(self, document: ProfessionalDocument) -> List:
        """Create table of contents for document."""
        story = []
        
        story.append(Paragraph("Table of Contents", getSampleStyleSheet()['Heading2']))
        story.append(Spacer(1, 12))
        
        toc_data = []
        for i, section in enumerate(document.sections, 1):
            toc_data.append([f"{i}.", section.title, f"Page {i + 2}"])  # Approximate page numbers
        
        toc_table = Table(toc_data, colWidths=[0.5*inch, 4*inch, 1*inch])
        toc_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        story.append(toc_table)
        story.append(Spacer(1, 20))
        
        return story
    
    async def _enhance_content_formatting(self, content: str) -> str:
        """Enhance content formatting for better presentation."""
        
        # Convert markdown-style formatting to HTML
        enhanced_content = content
        
        # Bold text
        enhanced_content = enhanced_content.replace('**', '<b>').replace('**', '</b>')
        enhanced_content = enhanced_content.replace('__', '<b>').replace('__', '</b>')
        
        # Italic text
        enhanced_content = enhanced_content.replace('*', '<i>').replace('*', '</i>')
        enhanced_content = enhanced_content.replace('_', '<i>').replace('_', '</i>')
        
        # Bullet points
        lines = enhanced_content.split('\n')
        enhanced_lines = []
        for line in lines:
            if line.strip().startswith('- ') or line.strip().startswith('• '):
                enhanced_lines.append(f"• {line.strip()[2:]}")
            elif line.strip().startswith('1. ') or line.strip().startswith('2. ') or line.strip().startswith('3. '):
                enhanced_lines.append(line.strip())
            else:
                enhanced_lines.append(line)
        
        enhanced_content = '\n'.join(enhanced_lines)
        
        return enhanced_content
    
    async def _create_signature_section(self, document: ProfessionalDocument) -> List:
        """Create signature section for document."""
        story = []
        
        signature = document.metadata.get("signature", {})
        
        story.append(Spacer(1, 30))
        story.append(Paragraph("Signature", getSampleStyleSheet()['Heading3']))
        story.append(Spacer(1, 12))
        
        if signature.get("type") == SignatureType.TEXT:
            story.append(Paragraph(f"<b>{signature.get('text', 'Digitally Signed')}</b>", getSampleStyleSheet()['Normal']))
        
        if signature.get("include_date"):
            story.append(Paragraph(f"Date: {signature.get('signature_date', datetime.now().strftime('%Y-%m-%d'))}", getSampleStyleSheet()['Normal']))
        
        return story
    
    def _add_watermark(self, canvas, doc):
        """Add watermark to PDF pages."""
        # This would be implemented with proper watermark functionality
        pass
    
    async def _export_to_advanced_word(
        self,
        document: ProfessionalDocument,
        request: DocumentExportRequest,
        watermark_config: Optional[Dict[str, Any]],
        signature_config: Optional[Dict[str, Any]],
        branding_config: Optional[Dict[str, Any]]
    ) -> str:
        """Export document to Word with advanced features."""
        
        filename = request.custom_filename or f"{document.title.replace(' ', '_')}_advanced.docx"
        file_path = self.output_dir / filename
        
        # Create Word document
        doc = DocxDocument()
        
        # Set document properties
        doc.core_properties.title = document.title
        if document.author:
            doc.core_properties.author = document.author
        doc.core_properties.created = document.date_created
        
        # Add branding header if configured
        if branding_config and document.metadata.get("branding"):
            await self._add_word_branding_header(doc, document)
        
        # Title
        title = doc.add_heading(document.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        if document.subtitle:
            subtitle = doc.add_heading(document.subtitle, 1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Author and company info
        if document.author or document.company:
            info_para = doc.add_paragraph()
            if document.author:
                info_para.add_run(f"Author: {document.author}\n").bold = True
            if document.company:
                info_para.add_run(f"Company: {document.company}\n").bold = True
            info_para.add_run(f"Date: {document.date_created.strftime('%B %d, %Y')}").bold = True
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page break
        doc.add_page_break()
        
        # Add table of contents if document is long
        if len(document.sections) > 5:
            await self._add_word_table_of_contents(doc, document)
        
        # Sections with enhanced formatting
        for section in document.sections:
            # Add heading
            heading = doc.add_heading(section.title, section.level)
            
            # Add content with enhanced formatting
            content_para = doc.add_paragraph()
            await self._add_enhanced_word_content(content_para, section.content, document.style)
        
        # Add signature if configured
        if signature_config and document.metadata.get("signature"):
            await self._add_word_signature(doc, document)
        
        # Save document
        doc.save(str(file_path))
        
        return str(file_path)
    
    async def _add_word_branding_header(self, doc: DocxDocument, document: ProfessionalDocument):
        """Add branding header to Word document."""
        
        branding = document.metadata.get("branding", {})
        
        if branding.get("company_name"):
            header_para = doc.add_paragraph()
            header_run = header_para.add_run(branding["company_name"])
            header_run.bold = True
            header_run.font.size = Pt(16)
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if branding.get("company_address"):
            address_para = doc.add_paragraph(branding["company_address"])
            address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add horizontal line
        doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Empty line
    
    async def _add_word_table_of_contents(self, doc: DocxDocument, document: ProfessionalDocument):
        """Add table of contents to Word document."""
        
        toc_heading = doc.add_heading("Table of Contents", 1)
        
        # Create table for TOC
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "No."
        hdr_cells[1].text = "Section"
        hdr_cells[2].text = "Page"
        
        # Add sections
        for i, section in enumerate(document.sections, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = section.title
            row_cells[2].text = str(i + 2)  # Approximate page numbers
        
        doc.add_page_break()
    
    async def _add_enhanced_word_content(self, para, content: str, style: DocumentStyle):
        """Add enhanced content to Word paragraph."""
        
        # Apply formatting based on content
        lines = content.split('\n')
        
        for line in lines:
            if line.strip().startswith('• ') or line.strip().startswith('- '):
                # Bullet point
                run = para.add_run(f"• {line.strip()[2:]}\n")
            elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
                # Numbered list
                run = para.add_run(f"{line.strip()}\n")
            elif line.strip().startswith('**') and line.strip().endswith('**'):
                # Bold text
                run = para.add_run(f"{line.strip()[2:-2]}\n")
                run.bold = True
            else:
                # Regular text
                run = para.add_run(f"{line}\n")
            
            # Apply style
            run.font.name = style.font_family
            run.font.size = Pt(style.font_size)
    
    async def _add_word_signature(self, doc: DocxDocument, document: ProfessionalDocument):
        """Add signature section to Word document."""
        
        signature = document.metadata.get("signature", {})
        
        doc.add_paragraph()  # Empty line
        doc.add_paragraph()  # Empty line
        
        sig_heading = doc.add_heading("Signature", 3)
        
        if signature.get("type") == SignatureType.TEXT:
            sig_para = doc.add_paragraph()
            sig_run = sig_para.add_run(signature.get("text", "Digitally Signed"))
            sig_run.bold = True
        
        if signature.get("include_date"):
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Date: {signature.get('signature_date', datetime.now().strftime('%Y-%m-%d'))}")
            date_run.bold = True
    
    async def _export_to_advanced_html(
        self,
        document: ProfessionalDocument,
        request: DocumentExportRequest,
        interactive_features: Optional[Dict[str, Any]]
    ) -> str:
        """Export document to HTML with advanced interactive features."""
        
        filename = request.custom_filename or f"{document.title.replace(' ', '_')}_advanced.html"
        file_path = self.output_dir / filename
        
        # Get interactive features
        features = document.metadata.get("interactive_features", {}) if interactive_features else {}
        
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{document.title}</title>
    <style>
        body {{
            font-family: {document.style.font_family};
            font-size: {document.style.font_size}px;
            line-height: {document.style.line_spacing};
            color: {document.style.body_color};
            background-color: {document.style.background_color};
            margin: {document.style.margin_top}in {document.style.margin_right}in {document.style.margin_bottom}in {document.style.margin_left}in;
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: {document.style.header_color};
        }}
        .title {{
            text-align: center;
            margin-bottom: 2em;
            border-bottom: 2px solid {document.style.accent_color};
            padding-bottom: 20px;
        }}
        .metadata {{
            text-align: center;
            margin-bottom: 2em;
            font-style: italic;
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
        }}
        .section {{
            margin-bottom: 2em;
            padding: 20px;
            border-left: 4px solid {document.style.accent_color};
            background-color: #ffffff;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .toc {{
            background-color: #f8f9fa;
            padding: 20px;
            border-radius: 5px;
            margin-bottom: 2em;
        }}
        .toc ul {{
            list-style-type: none;
            padding-left: 0;
        }}
        .toc li {{
            padding: 5px 0;
            border-bottom: 1px solid #dee2e6;
        }}
        .toc a {{
            text-decoration: none;
            color: {document.style.header_color};
        }}
        .toc a:hover {{
            color: {document.style.accent_color};
        }}
        .signature {{
            margin-top: 3em;
            padding: 20px;
            border-top: 2px solid {document.style.accent_color};
            text-align: right;
        }}
        .highlight {{
            background-color: #fff3cd;
            padding: 2px 4px;
            border-radius: 3px;
        }}
        .search-box {{
            position: fixed;
            top: 20px;
            right: 20px;
            background: white;
            padding: 10px;
            border: 1px solid #ccc;
            border-radius: 5px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .search-box input {{
            border: none;
            outline: none;
            padding: 5px;
        }}
    </style>
    {self._get_interactive_scripts(features)}
</head>
<body>
    {self._get_search_box(features)}
    
    <div class="title">
        <h1>{document.title}</h1>
        {f'<h2>{document.subtitle}</h2>' if document.subtitle else ''}
    </div>
    
    <div class="metadata">
        {f'<p><strong>Author:</strong> {document.author}</p>' if document.author else ''}
        {f'<p><strong>Company:</strong> {document.company}</p>' if document.company else ''}
        <p><strong>Date:</strong> {document.date_created.strftime('%B %d, %Y')}</p>
    </div>
    
    {self._get_table_of_contents(document, features)}
    
    <div class="content">
"""
        
        # Add sections with enhanced formatting
        for section in document.sections:
            heading_tag = f"h{min(section.level + 1, 6)}"
            enhanced_content = await self._enhance_html_content(section.content, features)
            
            html_content += f"""
        <div class="section" id="section-{section.order}">
            <{heading_tag}>{section.title}</{heading_tag}>
            <div class="section-content">{enhanced_content}</div>
        </div>
"""
        
        # Add signature if configured
        if document.metadata.get("signature"):
            html_content += self._get_html_signature(document)
        
        html_content += """
    </div>
</body>
</html>
"""
        
        # Write to file
        async with aiofiles.open(file_path, 'w', encoding='utf-8') as f:
            await f.write(html_content)
        
        return str(file_path)
    
    def _get_interactive_scripts(self, features: Dict[str, Any]) -> str:
        """Get JavaScript for interactive features."""
        
        if not features.get("search", True):
            return ""
        
        return """
    <script>
        function searchDocument() {
            const searchTerm = document.getElementById('searchInput').value.toLowerCase();
            const sections = document.querySelectorAll('.section');
            
            sections.forEach(section => {
                const content = section.textContent.toLowerCase();
                if (content.includes(searchTerm)) {
                    section.style.backgroundColor = '#fff3cd';
                    section.scrollIntoView({ behavior: 'smooth' });
                } else {
                    section.style.backgroundColor = '#ffffff';
                }
            });
        }
        
        function highlightText() {
            const searchTerm = document.getElementById('searchInput').value;
            if (searchTerm) {
                const walker = document.createTreeWalker(
                    document.querySelector('.content'),
                    NodeFilter.SHOW_TEXT,
                    null,
                    false
                );
                
                const textNodes = [];
                let node;
                while (node = walker.nextNode()) {
                    textNodes.push(node);
                }
                
                textNodes.forEach(textNode => {
                    const text = textNode.textContent;
                    const regex = new RegExp(`(${searchTerm})`, 'gi');
                    if (regex.test(text)) {
                        const highlightedText = text.replace(regex, '<span class="highlight">$1</span>');
                        const wrapper = document.createElement('div');
                        wrapper.innerHTML = highlightedText;
                        textNode.parentNode.replaceChild(wrapper, textNode);
                    }
                });
            }
        }
    </script>
"""
    
    def _get_search_box(self, features: Dict[str, Any]) -> str:
        """Get search box HTML."""
        
        if not features.get("search", True):
            return ""
        
        return """
    <div class="search-box">
        <input type="text" id="searchInput" placeholder="Search document..." onkeyup="searchDocument(); highlightText();">
    </div>
"""
    
    def _get_table_of_contents(self, document: ProfessionalDocument, features: Dict[str, Any]) -> str:
        """Get table of contents HTML."""
        
        if not features.get("bookmarks", True) or len(document.sections) <= 3:
            return ""
        
        toc_html = """
    <div class="toc">
        <h3>Table of Contents</h3>
        <ul>
"""
        
        for i, section in enumerate(document.sections):
            toc_html += f'            <li><a href="#section-{section.order}">{i+1}. {section.title}</a></li>\n'
        
        toc_html += """
        </ul>
    </div>
"""
        
        return toc_html
    
    async def _enhance_html_content(self, content: str, features: Dict[str, Any]) -> str:
        """Enhance HTML content with interactive features."""
        
        enhanced_content = content
        
        # Convert markdown-style formatting to HTML
        enhanced_content = enhanced_content.replace('**', '<strong>').replace('**', '</strong>')
        enhanced_content = enhanced_content.replace('*', '<em>').replace('*', '</em>')
        
        # Convert line breaks
        enhanced_content = enhanced_content.replace('\n', '<br>')
        
        # Add hyperlinks if enabled
        if features.get("hyperlinks", True):
            # Simple URL detection and conversion
            import re
            url_pattern = r'(https?://[^\s]+)'
            enhanced_content = re.sub(url_pattern, r'<a href="\1" target="_blank">\1</a>', enhanced_content)
        
        return enhanced_content
    
    def _get_html_signature(self, document: ProfessionalDocument) -> str:
        """Get HTML signature section."""
        
        signature = document.metadata.get("signature", {})
        
        if not signature:
            return ""
        
        sig_html = """
    <div class="signature">
        <h3>Signature</h3>
"""
        
        if signature.get("type") == SignatureType.TEXT:
            sig_html += f'        <p><strong>{signature.get("text", "Digitally Signed")}</strong></p>\n'
        
        if signature.get("include_date"):
            sig_html += f'        <p><strong>Date:</strong> {signature.get("signature_date", datetime.now().strftime("%Y-%m-%d"))}</p>\n'
        
        sig_html += "    </div>\n"
        
        return sig_html
    
    async def _export_document_standard(self, document: ProfessionalDocument, request: DocumentExportRequest) -> str:
        """Fall back to standard export if advanced features fail."""
        
        # Use parent class standard export
        response = await super().export_document(document, request)
        return response.file_path if response.success else None



























